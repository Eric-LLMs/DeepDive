"""Query Repository — multi-source import (file / learning / chat) behavior.

Covers the source-aware recall predicates (LEFT JOIN + file-only READY/domain), the shared
``write_query_repo_chunks`` helper, the chat segmentation helpers (default grouping + LLM
fallback), and the .docx / .pdf text-extraction dispatch. Uses scripted fakes — no database.
"""
import asyncio
from types import SimpleNamespace

from apps.worker.tasks import (
    _default_chat_pairs,
    _segment_chat,
    _strip_code_fence,
)
from core.infrastructure.ingest import (
    Chunk,
    extract_document_text,
    extract_text,
    write_query_repo_chunks,
)
from rag.recall.keyword import KeywordRecaller


# ── source-aware recall predicates (scripted fake session, no DB) ──
class _FakeSession:
    def __init__(self, rows):
        self.rows = rows
        self.sql = None
        self.params = None

    async def __aenter__(self):
        return self

    async def __aexit__(self, *exc):
        return False

    async def execute(self, stmt, params):
        self.sql = str(stmt)
        self.params = params
        return self

    def all(self):
        return self.rows


class _FakeSessionFactory:
    def __init__(self, rows=None):
        self.rows = rows or [SimpleNamespace(id="c1", content_en="text", meta={}, score=0.7)]
        self.session = None

    def __call__(self):
        self.session = _FakeSession(self.rows)
        return self.session


def _kw_sql(query="attention", filters=None):
    factory = _FakeSessionFactory()
    asyncio.run(KeywordRecaller(factory).recall(query, None, 5, filters))
    return factory.session.sql, factory.session.params


def test_recall_left_joins_assets_for_non_file_chunks():
    sql, _ = _kw_sql()
    assert "LEFT JOIN assets a ON a.id = c.asset_id" in sql


def test_recall_readiness_applies_only_to_file_chunks():
    sql, _ = _kw_sql("attention", {"user_id": "u1"})
    assert "c.asset_id IS NULL OR a.file_status = 'READY'" in sql


def test_recall_domain_filter_is_file_only():
    sql, params = _kw_sql("attention", {"user_id": "u1", "domain_id": "d1"})
    assert "c.asset_id IS NOT NULL AND a.domain_id = :domain_id::uuid" in sql
    assert params["domain_id"] == "d1"


def test_recall_still_leaf_only():
    sql, _ = _kw_sql()
    assert "c.chunk_kind = 'leaf'" in sql


# ── write_query_repo_chunks: embeds + bulk-inserts with source columns ──
class _FakeEmbedder:
    def __init__(self):
        self.calls = []

    async def embed(self, texts):
        self.calls.append(texts)
        return [[0.1] * 8 for _ in texts]


class _FakeChunkRepo:
    def __init__(self):
        self.bulk_calls = []

    async def bulk_insert(
        self, asset_id, user_id, workspace_id, chunks, source_type="file", source_id=None
    ):
        self.bulk_calls.append((asset_id, user_id, workspace_id, source_type, source_id, chunks))


def test_write_query_repo_chunks_persists_source_type_and_id(monkeypatch):
    from core.infrastructure import drive_repositories

    repo = _FakeChunkRepo()
    monkeypatch.setattr(drive_repositories, "SqlChunkRepository", lambda factory: repo)

    class _Factory:
        pass

    res = asyncio.run(
        write_query_repo_chunks(
            _Factory(),
            _FakeEmbedder(),
            chunks=[Chunk(content_en="hello world")],
            user_id="u1",
            source_type="chat",
            source_id="msg-1",
        )
    )

    assert res == {"chunks": 1}
    asset_id, user_id, workspace_id, source_type, source_id, rows = repo.bulk_calls[0]
    assert asset_id is None
    assert user_id == "u1"
    assert workspace_id is None
    assert source_type == "chat"
    assert source_id == "msg-1"
    assert rows[0]["content_en"] == "hello world"
    assert rows[0]["embedding"] == [0.1] * 8


# ── chat segmentation helpers ──
def _msg(role, text):
    return SimpleNamespace(role=role, text=text)


def test_default_chat_pairs_merge_follow_ups_into_one_answer():
    messages = [
        _msg("user", "What is RRF?"),
        _msg("assistant", "Reciprocal Rank Fusion."),
        _msg("assistant", "It fuses ranks."),
    ]
    pairs = _default_chat_pairs(messages)
    assert pairs == [
        {
            "question": "What is RRF?",
            "answer": "Reciprocal Rank Fusion.\nIt fuses ranks.",
            "indices": [0],
        }
    ]


def test_default_chat_pairs_splits_distinct_questions():
    messages = [
        _msg("user", "Q1"),
        _msg("assistant", "A1"),
        _msg("user", "Q2"),
        _msg("assistant", "A2"),
    ]
    pairs = _default_chat_pairs(messages)
    assert pairs == [
        {"question": "Q1", "answer": "A1", "indices": [0]},
        {"question": "Q2", "answer": "A2", "indices": [2]},
    ]


def test_strip_code_fence():
    assert _strip_code_fence('```json\n{"a": 1}\n```') == '{"a": 1}'


def test_segment_chat_parses_llm_json():
    class _GoodLLM:
        async def complete(self, prompt, system):
            # No indices field (older model) → fallback assigns user lines in order.
            return '[{"question": "Q1", "answer": "A1"}, {"question": "Q2", "answer": "A2"}]'

    messages = [_msg("user", "Q1"), _msg("assistant", "A1"), _msg("user", "Q2"), _msg("assistant", "A2")]
    pairs = asyncio.run(_segment_chat(messages, _GoodLLM()))
    assert pairs == [
        {"question": "Q1", "answer": "A1", "indices": [0]},
        {"question": "Q2", "answer": "A2", "indices": [2]},
    ]


def test_segment_chat_keeps_llm_merged_indices():
    class _MergingLLM:
        async def complete(self, prompt, system):
            # The LLM merged two follow-up user turns (indexes 2, 4) into one entry.
            return (
                '[{"question": "What is RRF?", "answer": "A1", "indices": [0]},'
                '{"question": "Explain more", "answer": "A2", "indices": [2, 4]}]'
            )

    messages = [
        _msg("user", "What is RRF?"),
        _msg("assistant", "A1"),
        _msg("user", "Wait, clarify"),
        _msg("assistant", "A2"),
        _msg("user", "And an example?"),
        _msg("assistant", "A3"),
    ]
    pairs = asyncio.run(_segment_chat(messages, _MergingLLM()))
    assert pairs == [
        {"question": "What is RRF?", "answer": "A1", "indices": [0]},
        {"question": "Explain more", "answer": "A2", "indices": [2, 4]},
    ]


def test_segment_chat_degrades_to_default_grouping_on_llm_failure():
    class _BoomLLM:
        async def complete(self, prompt, system):
            raise RuntimeError("model down")

    messages = [_msg("user", "Q"), _msg("assistant", "A")]
    pairs = asyncio.run(_segment_chat(messages, _BoomLLM()))
    assert pairs == [{"question": "Q", "answer": "A", "indices": [0]}]


# ── chat_session_import: incremental, flag-driven (no delete-and-rebuild) ──
class _Msg:
    def __init__(self, id, role, text, imported_rag=False):
        self.id = id
        self.role = role
        self.text = text
        self.imported_rag = imported_rag


class _RebuildResult:
    def __init__(self, messages, existing_ids):
        self._messages = messages
        self._existing_ids = existing_ids

    def scalars(self):
        return _RebuildScalars(self._messages)

    def all(self):
        return self._existing_ids


class _RebuildScalars:
    def __init__(self, rows):
        self._rows = rows

    def all(self):
        return self._rows


class _RebuildSession:
    def __init__(self, messages, existing_ids):
        self.messages = messages
        self.existing_ids = existing_ids
        self.committed = []

    async def __aenter__(self):
        return self

    async def __aexit__(self, *exc):
        return False

    async def execute(self, stmt, params=None):
        return _RebuildResult(self.messages, self.existing_ids)

    async def commit(self):
        self.committed.append(True)


class _RebuildFactory:
    def __init__(self, messages, existing_ids):
        self.messages = messages
        self.existing_ids = existing_ids
        self.session = None

    def __call__(self):
        self.session = _RebuildSession(self.messages, self.existing_ids)
        return self.session


class _RebuildStore:
    def __init__(self):
        self.events = []

    async def mark_running(self, uid):
        self.events.append(("running", uid))

    async def mark_succeeded(self, uid, result):
        self.events.append(("succeeded", uid, result))

    async def mark_failed(self, uid, error):
        self.events.append(("failed", uid, error))


class _RebuildRepo:
    def __init__(self):
        self.deleted = []

    async def delete_by_source(self, source_type, ids):
        self.deleted.append((source_type, list(ids)))


def _run_session_import(monkeypatch, messages, existing_ids):
    """Drive chat_session_import end-to-end with scripted segment/build/write fakes."""
    from rag import config_store

    import apps.worker.tasks as tasks_mod

    async def fake_load(session_factory):
        return {"pipeline": "test"}

    async def fake_segment(messages, llm):
        pairs = []
        for i, m in enumerate(messages):
            if m.role != "user":
                continue
            j = i + 1
            answer = messages[j].text if j < len(messages) else ""
            if answer:
                pairs.append({"question": m.text, "answer": answer, "indices": [i]})
        return pairs

    async def fake_build(text, cfg, doc_title=None, llm=None):
        return [Chunk(content_en=f"{doc_title}:{text}")]

    writes = []

    async def fake_write(session_factory, embedder, *, chunks, user_id, source_type, source_id):
        writes.append((source_id, user_id, source_type, chunks[0].content_en, chunks[0].meta))
        return {"chunks": len(chunks)}

    repo = _RebuildRepo()
    bumped = []

    async def fake_bump(redis):
        bumped.append(redis)

    monkeypatch.setattr(config_store, "load_config", fake_load)
    monkeypatch.setattr(tasks_mod, "_segment_chat", fake_segment)
    monkeypatch.setattr(tasks_mod, "build_chunks", fake_build)
    monkeypatch.setattr(tasks_mod, "write_query_repo_chunks", fake_write)
    monkeypatch.setattr(tasks_mod, "SqlChunkRepository", lambda sf: repo)
    monkeypatch.setattr(tasks_mod, "bump_corpus_version", fake_bump)

    factory = _RebuildFactory(messages, existing_ids)
    ctx = {
        "session_factory": factory,
        "job_store": _RebuildStore(),
        "llm": SimpleNamespace(),
        "embedder": SimpleNamespace(),
        "redis": SimpleNamespace(),
    }
    res = asyncio.run(
        tasks_mod.chat_session_import(
            ctx,
            "33333333-3333-3333-3333-333333333333",
            {"user_id": "22222222-2222-2222-2222-222222222222", "session_id": SID},
        )
    )
    return res, writes, repo.deleted, bumped, factory


SID = "11111111-1111-1111-1111-111111111111"


def test_session_import_skips_flagged_pairs_imports_only_new(monkeypatch):
    # 4 Q&A turns; the first two are already imported (flagged, flag-era chunks), the last
    # two are new appends. The import must skip the flagged pairs entirely — never re-embed
    # already-imported content — and only embed the new ones under their question's id.
    messages = [
        _Msg("m0", "user", "Q0", imported_rag=True),
        _Msg("m1", "assistant", "A0", imported_rag=True),
        _Msg("m2", "user", "Q1", imported_rag=True),
        _Msg("m3", "assistant", "A1", imported_rag=True),
        _Msg("m4", "user", "Q2", imported_rag=False),
        _Msg("m5", "assistant", "A2", imported_rag=False),
        _Msg("m6", "user", "Q3", imported_rag=False),
        _Msg("m7", "assistant", "A3", imported_rag=False),
    ]
    existing = []
    res, writes, deleted, bumped, factory = _run_session_import(monkeypatch, messages, existing)

    assert res == {"chunks": 2, "groups": 4}
    # Only the two unflagged pairs are written, keyed by their question's message id.
    assert [w[0] for w in writes] == ["m4", "m6"]
    # Per-pair idempotent delete-by-source (nothing existed → no-op) + one flag commit.
    assert deleted == [("chat", ["m4"]), ("chat", ["m6"])]
    assert bumped
    assert factory.session.committed == [True]
    assert writes[0][4]["covered"] == ["m4"]
    assert writes[0][4]["session_id"] == SID


def test_session_import_converts_legacy_session_on_first_reimport(monkeypatch):
    # A pre-flag whole-session import left positional keys (<session_id>:<i>) and no flags.
    # Re-importing purges those legacy keys once and re-embeds every pair under stable
    # message-id keys, so the session moves to the flag model without duplicating content.
    messages = [
        _Msg("m0", "user", "Q0"),
        _Msg("m1", "assistant", "A0"),
        _Msg("m2", "user", "Q1"),
        _Msg("m3", "assistant", "A1"),
    ]
    existing = [(f"{SID}:0",), (f"{SID}:1",)]
    res, writes, deleted, _, factory = _run_session_import(monkeypatch, messages, existing)

    assert res == {"chunks": 2, "groups": 2}
    # Legacy keys purged first, then each pair's old key replaced (idempotent).
    assert deleted == [("chat", [f"{SID}:0", f"{SID}:1"]), ("chat", ["m0"]), ("chat", ["m2"])]
    assert [w[0] for w in writes] == ["m0", "m2"]
    assert factory.session.committed == [True]


def test_session_import_reimports_after_answer_regenerated(monkeypatch):
    # Q0 was imported (question flagged), then its answer was regenerated as a fresh
    # assistant message id. The span now holds an unflagged assistant, so the pair re-imports
    # (replace the old chunk keyed by the question id) instead of being skipped.
    messages = [
        _Msg("m0", "user", "Q0", imported_rag=True),
        _Msg("m0a", "assistant", "A0 regenerated", imported_rag=False),
    ]
    existing = []
    res, writes, deleted, _, _ = _run_session_import(monkeypatch, messages, existing)

    assert res == {"chunks": 1, "groups": 1}
    assert deleted == [("chat", ["m0"])]
    assert [w[0] for w in writes] == ["m0"]
    assert "Q0" in writes[0][3] and "A0 regenerated" in writes[0][3]


def test_session_import_leaves_dangling_question_unflagged(monkeypatch):
    # A trailing user question with no answer is not part of any pair, so it is never flagged
    # (its content is not in the repo) and the all-imported run writes nothing.
    messages = [
        _Msg("m0", "user", "Q0", imported_rag=True),
        _Msg("m1", "assistant", "A0", imported_rag=True),
        _Msg("m2", "user", "Q1 dangling", imported_rag=False),
    ]
    existing = []
    res, writes, deleted, _, factory = _run_session_import(monkeypatch, messages, existing)

    assert res == {"chunks": 0, "groups": 1}
    assert writes == []
    assert deleted == []
    assert factory.session.committed == []  # nothing new to flag


def test_session_import_all_imported_fast_path_skips_llm(monkeypatch):
    # When every chat message is already flagged, the import must short-circuit BEFORE the
    # LLM segment call: re-importing a fully-imported session resolves instantly instead of
    # burning minutes in the model gateway and leaving the job "running".
    from rag import config_store

    import apps.worker.tasks as tasks_mod

    async def fake_load(session_factory):
        return {"pipeline": "test"}

    def boom(*args, **kwargs):
        raise AssertionError("fully-imported fast path must not call the LLM / embed / write")

    monkeypatch.setattr(config_store, "load_config", fake_load)
    monkeypatch.setattr(tasks_mod, "_segment_chat", boom)
    monkeypatch.setattr(tasks_mod, "build_chunks", boom)
    monkeypatch.setattr(tasks_mod, "write_query_repo_chunks", boom)
    monkeypatch.setattr(tasks_mod, "SqlChunkRepository", lambda sf: _RebuildRepo())
    monkeypatch.setattr(tasks_mod, "bump_corpus_version", boom)

    messages = [
        _Msg("m0", "user", "Q0", imported_rag=True),
        _Msg("m1", "assistant", "A0", imported_rag=True),
    ]
    factory = _RebuildFactory(messages, [])
    ctx = {
        "session_factory": factory,
        "job_store": _RebuildStore(),
        "llm": SimpleNamespace(),
        "embedder": SimpleNamespace(),
        "redis": SimpleNamespace(),
    }
    res = asyncio.run(
        tasks_mod.chat_session_import(
            ctx,
            "33333333-3333-3333-3333-333333333333",
            {"user_id": "22222222-2222-2222-2222-222222222222", "session_id": SID},
        )
    )
    assert res == {"chunks": 0, "groups": 0}
    assert factory.session.committed == []  # no flag writes, no commit


def test_pair_spans_cover_question_through_next_user_line():
    from apps.worker.tasks import _pair_spans

    msgs = [
        _Msg("a", "user", "Q0"), _Msg("b", "assistant", "A0"),
        _Msg("c", "user", "Q1"), _Msg("d", "assistant", "A1"),
        _Msg("e", "user", "Q2 dangling"),  # no answer — must not fall inside a span
    ]
    pairs = [{"indices": [0]}, {"indices": [2]}]
    assert _pair_spans(pairs, msgs) == [(0, 2), (2, 4)]


def test_span_imported_requires_whole_span_flagged():
    from apps.worker.tasks import _span_imported

    msgs = [_Msg("a", "user", "Q"), _Msg("b", "assistant", "A")]
    assert _span_imported(msgs, (0, 2), {"a": True, "b": True}) is True
    # A regenerated (unflagged) answer means the pair is not fully imported.
    assert _span_imported(msgs, (0, 2), {"a": True, "b": False}) is False


# ── text-extraction dispatch: .docx and .pdf ──
def test_extract_docx_paragraphs_and_table_cells():
    import io

    import docx as docx_lib

    buf = io.BytesIO()
    doc = docx_lib.Document()
    doc.add_paragraph("Hello paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "a"
    table.rows[0].cells[1].text = "b"
    doc.save(buf)

    text = extract_text(buf.getvalue(), "notes.docx")
    assert "Hello paragraph" in text
    assert "a | b" in text


def test_extract_document_text_delegates_plain_extract_for_txt():
    text = asyncio.run(extract_document_text("hello world".encode(), "notes.txt", llm=None))
    assert text == "hello world"


def test_extract_document_text_routes_pdf_to_pdf_extractor(monkeypatch):
    import core.infrastructure.pdf as pdf_mod

    calls = {}

    async def fake_extract(content, llm):
        calls["content"] = content
        calls["llm"] = llm
        return "pdf text"

    monkeypatch.setattr(pdf_mod, "extract_pdf_document", fake_extract)

    text = asyncio.run(extract_document_text(b"%PDF-1.4", "book.pdf", llm="LLM"))
    assert text == "pdf text"
    assert calls == {"content": b"%PDF-1.4", "llm": "LLM"}
