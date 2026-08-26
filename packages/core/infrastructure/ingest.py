"""Text extraction + chunking for the asset RAG-ingestion pipeline.

V1 extractors: plain text (``.txt``/``.md``/markdown and friends) and subtitle files
(``.srt``/``.vtt``/``.lrc``). Unsupported extensions raise :class:`UnsupportedFileType` so
the worker can mark the asset ``FAILED`` instead of crashing the job.

Chunking is strategy-dispatchable (``fixed`` / ``paragraph`` / ``sentence`` /
``semantic``), plus two P1/P2 enrichments:

- ``contextualize_chunks`` — Anthropic-style: an LLM writes a short context prefix per
  chunk so both vector and keyword recall match the chunk's surroundings.
- ``split_hierarchy`` — parent/leaf (small-to-big): recall surfaces small leaf chunks and
  ``parent_expand`` swaps a hit for its parent's fuller text.
"""
from __future__ import annotations

import asyncio
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path

from core.config import settings
from core.infrastructure import media

_SUBTITLE_EXTS = {".srt", ".vtt", ".lrc"}
_TEXT_EXTS = {".txt", ".md", ".markdown", ".text", ".log", ".json", ".csv"}


class UnsupportedFileType(ValueError):
    """Raised when no text extractor exists for the uploaded file type."""


@dataclass
class Chunk:
    """A prepared RAG chunk carrying every column the worker persists."""

    content_en: str
    content_cn: str | None = None
    meta: dict = field(default_factory=dict)
    chunk_kind: str = "leaf"                     # "leaf" (recalled) | "parent"
    parent_chunk_id: str | None = None           # leaf → its parent chunk id
    content_search: str | None = None            # jieba-segmented CJK index text
    id: str | None = None                        # client-assigned id (parents need one
    #                                            #  before their leaves reference it)


def _decode(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def extract_text(content: bytes, name: str) -> str:
    """Return searchable plain text for a file's bytes, dispatching on extension."""
    ext = Path(name).suffix.lower()
    text = _decode(content)
    if ext == ".vtt":
        cues = media.parse_vtt_text(text)
    elif ext == ".lrc":
        cues = media.parse_lrc_text(text)
    elif ext == ".srt":
        cues = media.parse_srt_text(text)
    elif ext in _TEXT_EXTS:
        return text
    elif ext == ".docx":
        return _extract_docx(content)
    else:
        raise UnsupportedFileType(f"no text extractor for '{ext or name}'")
    return "\n".join(c.text for c in cues if c.text)


def _extract_docx(content: bytes) -> str:
    """Extract paragraphs + table cells from a .docx (``python-docx``)."""
    import io

    import docx as docx_lib

    doc = docx_lib.Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(p for p in parts if p)


async def extract_document_text(content: bytes, name: str, llm=None) -> str:
    """Async dispatch for the ingest worker: PDF (body + vision tables) vs plain extract.

    PDF extraction needs the async LLM (table images → text), so it cannot live in the
    sync ``extract_text``; every other extension delegates there unchanged.
    """
    ext = Path(name).suffix.lower()
    if ext == ".pdf":
        from core.infrastructure.pdf import extract_pdf_document

        return await extract_pdf_document(content, llm)
    return extract_text(content, name)


async def write_query_repo_chunks(
    session_factory,
    embedder,
    *,
    chunks: list[Chunk],
    user_id,
    source_type: str,
    source_id: str | None,
    workspace_id=None,
) -> dict:
    """Embed + persist a batch of :class:`Chunk` objects as a non-file query-repo source.

    Shared by the worker jobs (``learning_import`` / ``chat_session_import``) and the
    single-pair ``/chat/import`` endpoint. ``asset_id`` stays NULL; the owner ``user_id``
    plus ``source_type`` / ``source_id`` (article id, sentence id, chat session id / message
    pair) are stored so recall scoping and idempotent re-import work.
    """
    from core.infrastructure.drive_repositories import SqlChunkRepository

    embeddings: list[list[float]] = []
    for i in range(0, len(chunks), settings.embed_batch_size):
        batch = chunks[i : i + settings.embed_batch_size]
        embeddings.extend(await embedder.embed([c.content_en for c in batch]))
    rows = [
        {
            "id": c.id,
            "content_en": c.content_en,
            "content_cn": c.content_cn,
            "meta": c.meta,
            "embedding": emb,
            "chunk_kind": c.chunk_kind,
            "parent_chunk_id": c.parent_chunk_id,
            "content_search": c.content_search,
        }
        for c, emb in zip(chunks, embeddings)
    ]
    repo = SqlChunkRepository(session_factory)
    await repo.bulk_insert(
        None, user_id, workspace_id, rows, source_type=source_type, source_id=source_id
    )
    return {"chunks": len(chunks)}


def split_chunks(
    text: str,
    chunk_chars: int | None = None,
    overlap: int | None = None,
    strategy: str = "fixed",
) -> list[str]:
    """Split text into leaf chunk texts by the configured strategy.

    ``fixed`` windows characters with overlap (the pre-refactor behaviour); ``paragraph``
    and ``sentence`` merge semantic units up to ``chunk_chars``; ``semantic`` is a
    placeholder that currently degrades to ``fixed`` (embed-based breakpoints are
    deferred). The default stays ``fixed`` so existing callers are unchanged.
    """
    chunk_chars = chunk_chars or settings.ingest_chunk_chars
    overlap = overlap or settings.ingest_chunk_overlap
    if not text.strip():
        return []
    strategy = strategy or "fixed"
    if strategy == "semantic":
        strategy = "fixed"  # interface exists; embed-based breakpoints not implemented yet
    if strategy == "paragraph":
        return _split_paragraphs(text, chunk_chars, overlap)
    if strategy == "sentence":
        return _split_sentences(text, chunk_chars, overlap)
    return _split_fixed(text, chunk_chars, overlap)


def _split_fixed(text: str, chunk_chars: int, overlap: int) -> list[str]:
    text = " ".join(text.split())
    if not text:
        return []
    if len(text) <= chunk_chars:
        return [text]
    step = max(1, chunk_chars - overlap)
    return [
        text[i : i + chunk_chars]
        for i in range(0, len(text), step)
    ]


def _split_paragraphs(text: str, chunk_chars: int, overlap: int) -> list[str]:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    return _merge_units(paragraphs, chunk_chars, overlap, separator="\n\n")


_SENT_BOUNDARY = re.compile(r"(?<=[.!?。！？])\s+|(?<=[\n])")


def _split_sentences(text: str, chunk_chars: int, overlap: int) -> list[str]:
    sentences = [s.strip() for s in _SENT_BOUNDARY.split(text) if s.strip()]
    return _merge_units(sentences, chunk_chars, overlap, separator=" ")


def _merge_units(units: list[str], chunk_chars: int, overlap: int, separator: str) -> list[str]:
    """Merge semantic units into chunks no longer than ``chunk_chars``.

    An oversized unit is windowed down via ``_split_fixed``; consecutive units that fit
    are joined with ``separator``. No overlap is re-introduced between merged units —
    paragraph/sentence chunking favours clean boundaries over the sliding window.
    """
    chunks: list[str] = []
    current = ""
    for unit in units:
        if len(unit) > chunk_chars:
            if current:
                chunks.append(current)
                current = ""
            windows = _split_fixed(unit, chunk_chars, overlap)
            chunks.extend(windows[:-1])
            current = windows[-1]
            continue
        if not current:
            current = unit
        elif len(current) + len(separator) + len(unit) <= chunk_chars:
            current = current + separator + unit
        else:
            chunks.append(current)
            current = unit
    if current:
        chunks.append(current)
    return chunks


async def build_chunks(text: str, config, *, doc_title: str, llm) -> list[Chunk]:
    """Prepare RAG chunks from raw text under a runtime pipeline config.

    Composes the chosen chunking strategy with the P1/P2 enrichments:
    parent/child hierarchy → Anthropic-style context prefixes → CJK ``content_search``.
    ``config`` is a :class:`RagPipelineConfig` (duck-typed so this module does not import
    the rag package at module scope).
    """
    from rag.query.cjk import segment  # lazy: rag is a sibling package

    chunking = config.chunking
    if config.parent_child:
        chunks = split_hierarchy(text, chunking.chunk_chars, chunking.overlap)
    else:
        chunks = [
            Chunk(content_en=c)
            for c in split_chunks(
                text, chunking.chunk_chars, chunking.overlap, chunking.strategy
            )
        ]
    if config.contextual:
        chunks = await contextualize_chunks(chunks, doc_title, llm)
    if config.cjk:
        for c in chunks:
            if not c.content_search:
                c.content_search = segment(c.content_en)
    return chunks


async def contextualize_chunks(
    chunks: list[Chunk],
    doc_title: str,
    llm,
    max_concurrency: int = 8,
) -> list[Chunk]:
    """Anthropic-style: prefix each leaf chunk with an LLM-written context snippet.

    The prompt follows Contextual Retrieval: given the document title and the chunk,
    the LLM writes 1-3 short sentences about where the chunk sits in the document. The
    prefix is stored as ``meta["context"]`` and prepended to ``content_en`` so both the
    embedding and the keyword index see it (the original text stays in ``meta["raw"]``).

    Concurrency is capped (``asyncio.Semaphore``) and a per-chunk LLM failure keeps the
    raw chunk — enrichment must never stall or fail the ingest job.
    """
    sem = asyncio.Semaphore(max_concurrency)
    system = (
        "You are a document-indexing assistant. Given a document title and a text "
        "chunk, write 1-3 short sentences (50-100 words total) describing the context "
        "of the chunk within the document: what topic it covers and how it relates to "
        "the surrounding material. Reply with only the context text."
    )

    async def one(c: Chunk) -> Chunk:
        if c.chunk_kind != "leaf":
            return c  # parents are already large; prefix leaves only
        if not c.content_en.strip():
            return c
        prompt = (
            f"Document title: {doc_title}\n\nChunk text:\n{c.content_en[:1200]}\n\n"
            "Write the context prefix:"
        )
        async with sem:
            try:
                context = (await llm.complete(prompt, system)).strip()
            except Exception:  # noqa: BLE001 - enrich or fall back, never fail ingest
                return c
        if not context:
            return c
        raw = c.content_en
        c.meta = {**c.meta, "raw": raw, "context": context}
        c.content_en = f"{context}\n{raw}"
        return c

    return list(await asyncio.gather(*(one(c) for c in chunks)))


def split_hierarchy(
    text: str,
    chunk_chars: int = 1200,
    overlap: int = 150,
    parent_chars: int | None = None,
) -> list[Chunk]:
    """Small-to-big split: produce parent chunks + their leaf children.

    Leaf chunks (recalled) record ``parent_chunk_id``; parent chunks hold a larger window
    of text (default 3× the leaf size) so ``parent_expand`` can widen a hit's context.
    Parents carry ``chunk_kind="parent"`` and are NOT recalled directly.

    Assignment is range-based: every leaf is linked to the parent whose window contains
    the leaf's midpoint, guaranteeing each leaf has exactly one parent and siblings share
    a parent (which ``parent_expand`` dedups). Parents get client-assigned UUIDs so their
    leaves can reference them before the rows are inserted.
    """
    parent_chars = parent_chars or chunk_chars * 3
    text = " ".join(text.split())
    if not text:
        return []
    leaf_ranges = _fixed_ranges(text, chunk_chars, overlap)
    if len(leaf_ranges) <= 1:
        # Single leaf → it is its own context; no hierarchy needed.
        return [Chunk(content_en=text)] if leaf_ranges else []

    parent_ranges = _fixed_ranges(text, parent_chars, overlap)
    out: list[Chunk] = []
    parents: list[tuple[int, int, str]] = []
    for p_start, p_end in parent_ranges:
        pid = str(uuid.uuid4())
        parents.append((p_start, p_end, pid))
        out.append(
            Chunk(
                id=pid,
                content_en=text[p_start:p_end],
                chunk_kind="parent",
                meta={"chunk_kind": "parent", "span": [p_start, p_end]},
            )
        )

    for l_start, l_end in leaf_ranges:
        mid = (l_start + l_end) // 2
        parent_id = next((pid for p_start, p_end, pid in parents if p_start <= mid < p_end), None)
        out.append(
            Chunk(
                content_en=text[l_start:l_end],
                chunk_kind="leaf",
                parent_chunk_id=parent_id,
            )
        )
    return out


def _fixed_ranges(text: str, chunk_chars: int, overlap: int) -> list[tuple[int, int]]:
    """Character ranges for the fixed-window split (used by the hierarchy builder)."""
    if len(text) <= chunk_chars:
        return [(0, len(text))]
    step = max(1, chunk_chars - overlap)
    return [(i, i + chunk_chars) for i in range(0, len(text), step)]
