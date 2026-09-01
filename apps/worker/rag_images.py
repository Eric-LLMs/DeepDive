"""Extract embedded images from a RAG-ingested document and save them as cloud-drive assets.

The agent's vision tool reads images by ``asset_id``, so a document's figures only become
usable during RAG retrieval if they are persisted as drive files AND referenced from the
chunks that discuss them. This module does the "save + reference" half:

- ``scan_embedded_images`` walks the document and returns, per anchor, the images embedded
  there. Anchors are 1-based page numbers for PDF (via PyMuPDF ``get_images`` /
  ``extract_image``) and paragraph indexes for DOCX (via the ``a:blip r:embed`` drawing
  anchors). DOCX has no native pages, so paragraph anchoring is the closest equivalent to
  "the image that travels with this chunk".
- ``save_images`` persists each image with :meth:`DriveService.save_artifact` into a
  dedicated ``RAG 图片/<doc>`` folder, records ``source_asset_id`` (the PDF/DOCX asset) on
  the image asset, and dedupes by ``(source_asset_id, content-hash)`` so re-ingesting the
  same document reuses existing image assets instead of duplicating rows.

Chunk association happens in ``apps/worker/tasks.py`` via ``build_chunks(on_split=...)``:
the ``[[PAGE:n]]`` / ``[[PARA:n]]`` markers in the extracted text are stripped there and
replaced with ``meta["pages"]`` / ``meta["paras"]`` + ``meta["image_ids"]`` (a deduped union
across every page/paragraph the chunk covers).
"""
from __future__ import annotations

import hashlib
import io
import re

_RASTER_EXTS = {"png", "jpg", "jpeg", "gif", "webp", "bmp"}
_EXT_MIME = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "webp": "image/webp",
    "bmp": "image/bmp",
}

# Characters that would turn a doc title into nested folders / broken names.
_FOLDER_BAD = re.compile(r'[\\/:*?"<>|]')


def _sanitize_folder(name: str) -> str:
    cleaned = _FOLDER_BAD.sub("_", (name or "doc").strip()).strip(". ")
    return cleaned or "doc"


def _scan_pdf(data: bytes) -> dict[int, list[dict]]:
    """Map 1-based page number → embedded raster images of that page."""
    import pymupdf

    out: dict[int, list[dict]] = {}
    doc = pymupdf.open(stream=data, filetype="pdf")
    try:
        for pno in range(doc.page_count):
            images: list[dict] = []
            for img in doc[pno].get_images(full=True):
                xref = img[0]
                info = doc.extract_image(xref)
                ext = (info.get("ext") or "").lower()
                if ext not in _RASTER_EXTS:
                    continue
                images.append(
                    {
                        "name": f"p{pno + 1}_{xref}.{ext}",
                        "mime": _EXT_MIME.get(ext, "image/png"),
                        "data": info["image"],
                    }
                )
            if images:
                out[pno + 1] = images
    finally:
        doc.close()
    return out


def _blip_rids(paragraph) -> list[str]:
    """Return the image relationship ids anchored in a docx paragraph (a:blip r:embed)."""
    from docx.oxml.ns import qn

    rids: list[str] = []
    for blip in paragraph._p.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            rids.append(rid)
    return rids


def _scan_docx(data: bytes) -> dict[int, list[dict]]:
    """Map paragraph index → images anchored in that paragraph.

    Headers/footers/text boxes/table cells are out of scope (they are not ``document.paragraphs``);
    the paragraph index matches ``_extract_docx(..., para_markers=True)`` exactly, since both
    iterate ``document.paragraphs`` in document order.
    """
    import docx as docx_lib

    doc = docx_lib.Document(io.BytesIO(data))
    part_info: dict[str, dict] = {}
    for rId, rel in doc.part.rels.items():
        if not (rel.reltype or "").endswith("/image"):
            continue
        part = rel.target_part
        name = str(part.partname)
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        if ext not in _RASTER_EXTS:
            continue
        stem = name.rsplit("/", 1)[-1] or f"{rId}.{ext}"
        part_info[rId] = {
            "mime": part.content_type or _EXT_MIME.get(ext, "image/png"),
            "data": part.blob,
            "stem": stem,
        }
    out: dict[int, list[dict]] = {}
    for i, para in enumerate(doc.paragraphs):
        images: list[dict] = []
        for rid in _blip_rids(para):
            info = part_info.get(rid)
            if info is None:
                continue
            images.append(
                {"name": f"p{i}_{info['stem']}", "mime": info["mime"], "data": info["data"]}
            )
        if images:
            out[i] = images
    return out


def scan_embedded_images(data: bytes, name: str) -> dict[int, list[dict]]:
    """Return anchor → images for a PDF/DOCX, or ``{}`` for formats with no image pass."""
    ext = (name or "").rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return _scan_pdf(data)
    if ext == "docx":
        return _scan_docx(data)
    return {}


async def save_images(
    scans: dict[int, list[dict]],
    doc_title: str,
    user_id,
    workspace_id,
    source_asset_id,
    drive,
) -> dict[int, list[str]]:
    """Persist scanned images as drive assets; return ``anchor → [asset_id, ...]``.

    Dedupes on ``(source_asset_id, content-hash)`` so a re-ingest of the same document
    reuses existing image assets (no duplicate rows), and so the same image anchored on
    multiple pages resolves to one asset referenced by all of them.
    """
    folder = f"RAG 图片/{_sanitize_folder(doc_title)}"
    out: dict[int, list[str]] = {}
    assets = drive.assets
    for key, images in scans.items():
        ids: list[str] = []
        for img in images:
            digest = hashlib.sha256(img["data"]).hexdigest()
            existing = await assets.get_by_source_content(source_asset_id, digest)
            if existing is not None:
                ids.append(str(existing.id))
                continue
            asset = await drive.save_artifact(
                user_id,
                img["name"],
                img["mime"],
                img["data"],
                folder_path=folder,
                workspace_id=workspace_id,
                source_asset_id=source_asset_id,
            )
            ids.append(str(asset.id))
        out[key] = ids
    return out
